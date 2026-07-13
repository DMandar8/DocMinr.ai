"""
DOCX Parser using python-docx
"""
from typing import Dict, Any
import docx
from .base import BaseParser

class DOCXParser(BaseParser):
    """Parser for DOCX documents"""
    
    def extract_text(self) -> str:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(self.file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            return self.clean_text(text)
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def extract_metadata(self) -> Dict[str, Any]:
        """Extract metadata from DOCX"""
        try:
            doc = docx.Document(self.file_path)
            text = self.extract_text()
            words = text.split()
            
            return {
                "page_count": 1,
                "paragraph_count": len(doc.paragraphs),
                "word_count": len(words),
                "char_count": len(text),
                "file_name": self.file_name,
                "file_size": self.file_path.stat().st_size,
            }
        except Exception as e:
            raise Exception(f"Metadata extraction failed: {str(e)}")